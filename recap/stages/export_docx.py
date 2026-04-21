"""Phase 4 slice: optional DOCX export.

Reads the same artifacts `recap assemble` and `recap export-html` read
(`job.json`, `metadata.json`, `transcript.json`, and — when present —
`selected_frames.json` + `chapter_candidates.json`) and writes
`report.docx` via `python-docx`. No Markdown parsing, no network, no
VLM/LLM calls. Images from `candidate_frames/` are embedded into the
DOCX package with `Document.add_picture(...)`; no image files are
copied, renamed, or rewritten on disk.

This stage is opt-in via `recap export-docx --job <path>`. It is NOT
invoked by `recap run` and it is NOT part of `job.STAGES`. It does not
modify `report.md`, `report.html`, `selected_frames.json`, or any
upstream artifact.

Validation follows the same selected-path contract as the Markdown and
HTML exports. The shared validators, formatters, and safety helpers
live in `recap/stages/report_helpers.py` so all three report stages
stay in lockstep.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches

from ..job import COMPLETED, FAILED, RUNNING, JobPaths, read_job, update_stage
from .report_helpers import (
    apply_frame_review_to_chapter as _apply_frame_review_to_chapter,
    caption_for as _caption_for,
    check_hero_coherence as _check_hero_coherence,
    check_supporting_coherence as _check_supporting_coherence,
    collapse_whitespace as _collapse_whitespace,
    format_ts as _format_ts,
    insights_chapters_by_index as _insights_chapters_by_index,
    iter_transcript_utterances as _iter_transcript_utterances,
    load_chapter_titles_overlay as _load_chapter_titles_overlay,
    load_frame_review_overlay as _load_frame_review_overlay,
    load_insights as _load_insights,
    load_speaker_names_overlay as _load_speaker_names_overlay,
    resolve_chapter_title as _resolve_chapter_title,
    resolve_speaker_label as _resolve_speaker_label,
    summarize_metadata as _summarize_metadata,
    validate_chapter_candidates as _validate_chapter_candidates,
    validate_selected_frames as _validate_selected_frames,
)


_IMAGE_WIDTH_INCHES = 6.0



def _add_caption(doc, caption: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.italic = True


def _render_overview(doc, insights: dict) -> None:
    overview = insights.get("overview") or {}
    doc.add_heading("Overview", level=2)
    short_summary = (overview.get("short_summary") or "").strip()
    if short_summary:
        doc.add_paragraph(short_summary)
    detailed = (overview.get("detailed_summary") or "").strip()
    if detailed and detailed != short_summary:
        doc.add_paragraph(detailed)

    bullets = overview.get("quick_bullets") or []
    if bullets:
        doc.add_heading("Quick bullets", level=3)
        for b in bullets:
            if isinstance(b, str) and b.strip():
                doc.add_paragraph(b.strip(), style="List Bullet")

    actions = insights.get("action_items") or []
    if actions:
        doc.add_heading("Action items", level=3)
        for ai in actions:
            if not isinstance(ai, dict):
                continue
            text = (ai.get("text") or "").strip()
            if not text:
                continue
            bits = [text]
            stamp = ai.get("timestamp_seconds")
            if isinstance(stamp, (int, float)) and not isinstance(stamp, bool):
                bits.append(f"[{_format_ts(float(stamp))}]")
            ch_idx = ai.get("chapter_index")
            if isinstance(ch_idx, int):
                bits.append(f"Chapter {ch_idx}")
            owner = ai.get("owner")
            if isinstance(owner, str) and owner.strip():
                bits.append(f"Owner: {owner.strip()}")
            due = ai.get("due")
            if isinstance(due, str) and due.strip():
                bits.append(f"Due: {due.strip()}")
            suffix = (" — " + " · ".join(bits[1:])) if len(bits) > 1 else ""
            doc.add_paragraph(
                f"☐ {bits[0]}{suffix}", style="List Bullet"
            )


def _add_chapter_insights_block(doc, ch: dict) -> None:
    summary = (ch.get("summary") or "").strip()
    if summary:
        doc.add_paragraph(summary)
    bullets = ch.get("bullets") or []
    for b in bullets:
        if isinstance(b, str) and b.strip():
            doc.add_paragraph(b.strip(), style="List Bullet")
    action_items = ch.get("action_items") or []
    if action_items:
        doc.add_paragraph("Action items:")
        for ai in action_items:
            if isinstance(ai, str) and ai.strip():
                doc.add_paragraph(f"☐ {ai.strip()}", style="List Bullet")
    speakers = ch.get("speaker_focus") or []
    if speakers:
        p = doc.add_paragraph()
        r = p.add_run(f"Speakers: {', '.join(speakers)}")
        r.italic = True


def _render_insights_only_chapters(
    doc,
    insights: dict,
    chapter_titles_overlay: dict[int, str] | None = None,
) -> None:
    doc.add_heading("Chapters", level=2)
    overlay = chapter_titles_overlay or {}
    for ch in insights.get("chapters") or []:
        if not isinstance(ch, dict):
            continue
        idx = ch.get("index")
        start = _format_ts(ch.get("start_seconds"))
        end = _format_ts(ch.get("end_seconds"))
        heading = f"Chapter {idx}"
        effective_title = None
        if isinstance(idx, int):
            effective_title = _resolve_chapter_title(
                idx,
                custom_by_idx=overlay,
                insights_title=(
                    ch.get("title")
                    if isinstance(ch.get("title"), str)
                    else None
                ),
            )
        if effective_title:
            heading += f" — {effective_title}"
        heading += f" [{start} – {end}]"
        doc.add_heading(heading, level=3)
        _add_chapter_insights_block(doc, ch)


def _add_image(doc, image_path: Path) -> None:
    doc.add_picture(str(image_path), width=Inches(_IMAGE_WIDTH_INCHES))



def _render_chapters(
    doc,
    selected: dict,
    chapter_text_by_index: dict[int, str],
    frames_dir: Path,
    insights_by_idx: dict[int, dict] | None = None,
    chapter_titles_overlay: dict[int, str] | None = None,
    frame_review_overlay: dict[str, dict] | None = None,
) -> None:
    doc.add_heading("Chapters", level=2)

    title_overlay = chapter_titles_overlay or {}
    review_overlay = frame_review_overlay or {}

    for ch in selected["chapters"]:
        ch_idx = ch["chapter_index"]
        if ch_idx not in chapter_text_by_index:
            raise RuntimeError(
                f"chapter_candidates.json has no chapter with index {ch_idx} "
                "required by selected_frames.json"
            )

        # Coherence checks run on the raw artifact first. The overlay
        # is then applied to the validated output.
        hero_frame_raw = _check_hero_coherence(ch)
        _check_supporting_coherence(ch)

        start = _format_ts(ch.get("start_seconds") or 0.0)
        end = _format_ts(ch.get("end_seconds") or 0.0)
        insight = (insights_by_idx or {}).get(ch_idx)
        effective_title = _resolve_chapter_title(
            ch_idx,
            custom_by_idx=title_overlay,
            insights_title=(insight or {}).get("title") if insight else None,
        )
        heading = f"Chapter {ch_idx}"
        if effective_title:
            heading += f" — {effective_title}"
        heading += f" [{start} – {end}]"
        doc.add_heading(heading, level=3)
        if insight is not None:
            _add_chapter_insights_block(doc, insight)

        effective_hero, effective_supporting = (
            _apply_frame_review_to_chapter(
                ch, hero_frame_raw, review_overlay,
            )
        )

        if effective_hero is not None:
            frame_file = effective_hero["frame_file"]
            image_path = frames_dir / frame_file
            if not image_path.exists():
                raise RuntimeError(
                    f"missing candidate frame: candidate_frames/{frame_file}"
                )
            _add_image(doc, image_path)
            caption = _caption_for(effective_hero)
            if caption:
                _add_caption(doc, caption)

        for support_frame in effective_supporting:
            frame_file = support_frame["frame_file"]
            image_path = frames_dir / frame_file
            if not image_path.exists():
                raise RuntimeError(
                    f"missing candidate frame: candidate_frames/{frame_file}"
                )
            _add_image(doc, image_path)
            caption = _caption_for(support_frame)
            if caption:
                _add_caption(doc, caption)

        body_text = _collapse_whitespace(chapter_text_by_index[ch_idx])
        if body_text:
            doc.add_paragraph(body_text)


def build_document(
    job: dict,
    meta_summary: dict,
    transcript: dict | None,
    selected: dict | None,
    chapter_text_by_index: dict[int, str] | None,
    frames_dir: Path,
    insights: dict | None = None,
    chapter_titles_overlay: dict[int, str] | None = None,
    speaker_names_overlay: dict[str, str] | None = None,
    frame_review_overlay: dict[str, dict] | None = None,
):
    doc = Document()
    title = job.get("original_filename") or job.get("job_id") or "Recap"
    doc.add_heading(f"Recap: {title}", level=1)

    if job.get("job_id"):
        doc.add_paragraph(f"Job ID: {job['job_id']}")
    if job.get("original_filename"):
        doc.add_paragraph(f"Source file: {job['original_filename']}")
    if job.get("created_at"):
        doc.add_paragraph(f"Created: {job['created_at']}")

    if insights is not None:
        _render_overview(doc, insights)

    if selected is None and insights is not None and (
        insights.get("chapters") or []
    ):
        _render_insights_only_chapters(
            doc,
            insights,
            chapter_titles_overlay=chapter_titles_overlay,
        )

    doc.add_heading("Media", level=2)
    dur = meta_summary.get("duration_seconds")
    if dur is not None:
        doc.add_paragraph(f"Duration: {_format_ts(dur)} ({dur:.2f}s)")
    if meta_summary.get("format_name"):
        doc.add_paragraph(f"Container: {meta_summary['format_name']}")
    v = meta_summary.get("video")
    if v:
        res = f"{v.get('width')}x{v.get('height')}" if v.get("width") else "unknown"
        doc.add_paragraph(
            f"Video: {v.get('codec')} {res} @ {v.get('frame_rate')}"
        )
    a = meta_summary.get("audio")
    if a:
        doc.add_paragraph(
            f"Audio: {a.get('codec')} {a.get('sample_rate')} Hz, "
            f"{a.get('channels')} ch"
        )

    if selected is not None and chapter_text_by_index is not None:
        insights_by_idx = (
            _insights_chapters_by_index(insights) if insights else None
        )
        _render_chapters(
            doc,
            selected,
            chapter_text_by_index,
            frames_dir,
            insights_by_idx=insights_by_idx,
            chapter_titles_overlay=chapter_titles_overlay,
            frame_review_overlay=frame_review_overlay,
        )

    doc.add_heading("Transcript", level=2)
    if transcript is None:
        doc.add_paragraph("No transcript available.")
        return doc

    engine = transcript.get("engine")
    model = transcript.get("model")
    doc.add_paragraph(f"Engine: {engine} (model {model})")
    if transcript.get("language"):
        doc.add_paragraph(f"Detected language: {transcript['language']}")

    utterances = _iter_transcript_utterances(transcript)
    speaker_labels = speaker_names_overlay or {}
    if utterances:
        doc.add_paragraph(f"Utterances: {len(utterances)}")
        doc.add_heading("Utterances", level=3)
        for u in utterances:
            if not isinstance(u, dict):
                continue
            start = _format_ts(u.get("start") or 0.0)
            end = _format_ts(u.get("end") or 0.0)
            text = (u.get("text") or "").strip()
            if not text:
                continue
            label = _resolve_speaker_label(
                u.get("speaker"), speaker_labels,
            )
            prefix = f"{label}: " if label else ""
            doc.add_paragraph(
                f"[{start} – {end}] {prefix}{text}", style="List Bullet"
            )
        return doc

    segs = transcript.get("segments", []) or []
    doc.add_paragraph(f"Segments: {len(segs)}")

    doc.add_heading("Segments", level=3)
    for seg in segs:
        start = _format_ts(seg.get("start") or 0.0)
        end = _format_ts(seg.get("end") or 0.0)
        text = (seg.get("text") or "").strip()
        if not text:
            continue
        doc.add_paragraph(
            f"[{start} – {end}] {text}", style="List Bullet"
        )

    return doc


def run(paths: JobPaths, force: bool = False) -> Path:
    update_stage(paths, "export_docx", RUNNING)
    tmp = paths.report_docx.with_suffix(".docx.tmp")
    try:
        if not force and paths.report_docx.exists():
            update_stage(
                paths, "export_docx", COMPLETED, extra={"skipped": True}
            )
            return paths.report_docx

        job = read_job(paths)

        meta_summary: dict = {}
        if paths.metadata_json.exists():
            with open(paths.metadata_json, "r", encoding="utf-8") as f:
                meta_summary = _summarize_metadata(json.load(f))

        transcript: dict | None = None
        if paths.transcript_json.exists():
            with open(paths.transcript_json, "r", encoding="utf-8") as f:
                transcript = json.load(f)

        insights = _load_insights(paths.insights_json)

        speaker_overlay = _load_speaker_names_overlay(
            paths.speaker_names_json,
        )
        chapter_title_overlay = _load_chapter_titles_overlay(
            paths.chapter_titles_json,
        )
        frame_review_overlay = _load_frame_review_overlay(
            paths.frame_review_json,
        )

        selected: dict | None = None
        chapter_text_by_index: dict[int, str] | None = None
        if paths.selected_frames_json.exists():
            try:
                with open(
                    paths.selected_frames_json, "r", encoding="utf-8"
                ) as f:
                    selected_raw = json.load(f)
            except json.JSONDecodeError as e:
                raise RuntimeError(
                    f"selected_frames.json malformed: invalid JSON: {e.msg}"
                ) from e
            selected = _validate_selected_frames(selected_raw)

            if not paths.chapter_candidates_json.exists():
                raise RuntimeError(
                    "chapter_candidates.json is required when "
                    "selected_frames.json is present but was not found"
                )
            try:
                with open(
                    paths.chapter_candidates_json, "r", encoding="utf-8"
                ) as f:
                    chapters_raw = json.load(f)
            except json.JSONDecodeError as e:
                raise RuntimeError(
                    f"chapter_candidates.json malformed: invalid JSON: {e.msg}"
                ) from e
            chapter_text_by_index = _validate_chapter_candidates(chapters_raw)

        doc = build_document(
            job,
            meta_summary,
            transcript,
            selected,
            chapter_text_by_index,
            paths.candidate_frames_dir,
            insights=insights,
            chapter_titles_overlay=chapter_title_overlay,
            speaker_names_overlay=speaker_overlay,
            frame_review_overlay=frame_review_overlay,
        )

        doc.save(str(tmp))
        tmp.replace(paths.report_docx)

        extra: dict = {
            "report": paths.report_docx.name,
            "bytes": paths.report_docx.stat().st_size,
            "embedded_selected_frames": selected is not None,
            "embedded_insights": insights is not None,
            "overlays": {
                "speaker_names": bool(speaker_overlay),
                "chapter_titles": bool(chapter_title_overlay),
                "frame_review": bool(frame_review_overlay),
            },
        }
        update_stage(paths, "export_docx", COMPLETED, extra=extra)
        return paths.report_docx
    except Exception as e:
        if tmp.exists():
            try:
                tmp.unlink()
            except OSError:
                pass
        update_stage(
            paths, "export_docx", FAILED, error=f"{type(e).__name__}: {e}"
        )
        raise
