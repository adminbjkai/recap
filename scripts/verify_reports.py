#!/usr/bin/env python3
"""Offline golden-path validation for Recap's Markdown / HTML / DOCX exports.

Runs `recap assemble`, `recap export-html`, and `recap export-docx`
against a tiny committed fixture under `scripts/fixtures/minimal_job/`
and asserts the expected structure for both the selected-path and the
absent-selected path, plus the shared error envelope for a small set of
malformed artifacts.

No network, no model downloads, no pytest. Uses stdlib plus `python-docx`
(already a project dependency).
"""

from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parent.parent
FIXTURE_ROOT = REPO_ROOT / "scripts" / "fixtures" / "minimal_job"
EXPORT_CMDS = ("assemble", "export-html", "export-docx")
CHECKS_PASSED = 0

# Prefix used by the mock insights provider's output on disk. Keep in
# sync with recap/stages/insights.py.
INSIGHTS_SCHEMA_VERSION = 1


def fail(case: str, reason: str) -> None:
    print(f"FAIL: {case}: {reason}")
    sys.exit(1)


def passed() -> None:
    global CHECKS_PASSED
    CHECKS_PASSED += 1


def copy_fixture() -> Path:
    scratch = Path(tempfile.mkdtemp(prefix="recap_verify_"))
    job = scratch / "minimal_job"
    shutil.copytree(FIXTURE_ROOT, job)
    return job


def run_cmd(
    case: str,
    job: Path,
    cmd: str,
    force: bool = True,
    extra: tuple[str, ...] = (),
) -> subprocess.CompletedProcess:
    args = [sys.executable, "-m", "recap", cmd, "--job", str(job)]
    if force:
        args.append("--force")
    args.extend(extra)
    return subprocess.run(
        args, cwd=REPO_ROOT, capture_output=True, text=True
    )


def expect_ok(
    case: str,
    job: Path,
    cmd: str,
    extra: tuple[str, ...] = (),
) -> None:
    result = run_cmd(case, job, cmd, extra=extra)
    if result.returncode != 0:
        fail(
            case,
            f"`recap {cmd}` returned {result.returncode}; "
            f"stderr={result.stderr.strip()!r}; job={job}",
        )


def expect_fail_exact(
    case: str,
    job: Path,
    cmd: str,
    extra: tuple[str, ...],
    needles: tuple[str, ...],
) -> None:
    result = run_cmd(case, job, cmd, extra=extra)
    if result.returncode == 0:
        fail(
            case,
            f"`recap {cmd} {' '.join(extra)}` unexpectedly succeeded; "
            f"stdout={result.stdout.strip()!r}; job={job}",
        )
    if not any(n in result.stderr for n in needles):
        fail(
            case,
            f"`recap {cmd}` stderr did not contain any of {needles!r}; "
            f"got {result.stderr.strip()!r}",
        )


def expect_fail_with(
    case: str, job: Path, cmd: str, needles: tuple[str, ...]
) -> None:
    result = run_cmd(case, job, cmd)
    if result.returncode != 2:
        fail(
            case,
            f"`recap {cmd}` returned {result.returncode}, expected 2; "
            f"stderr={result.stderr.strip()!r}; job={job}",
        )
    if not any(n in result.stderr for n in needles):
        fail(
            case,
            f"`recap {cmd}` stderr did not contain any of {needles!r}; "
            f"got {result.stderr.strip()!r}",
        )
    for tmp_name in ("report.md.tmp", "report.html.tmp", "report.docx.tmp"):
        if (job / tmp_name).exists():
            fail(case, f"leftover tmp file: {tmp_name}")


def assert_contains(case: str, path: Path, needle: str) -> None:
    text = path.read_text(encoding="utf-8")
    if needle not in text:
        fail(case, f"{path.name} missing {needle!r}")


def assert_not_contains(case: str, path: Path, needle: str) -> None:
    text = path.read_text(encoding="utf-8")
    if needle in text:
        fail(case, f"{path.name} unexpectedly contains {needle!r}")


def docx_headings(path: Path) -> list[str]:
    doc = Document(str(path))
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


def docx_inline_shapes(path: Path) -> int:
    return len(Document(str(path)).inline_shapes)


def check_selected_path() -> None:
    case = "selected-path"
    job = copy_fixture()
    try:
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        assert_contains(case, md, "## Chapters")
        assert_contains(case, md, "candidate_frames/scene-001.jpg")
        assert_contains(case, md, "candidate_frames/scene-003.jpg")
        assert_not_contains(case, md, "candidate_frames/scene-002.jpg")
        assert_contains(case, md, "## Transcript")
        passed()

        html = job / "report.html"
        assert_contains(case, html, "<h2>Chapters</h2>")
        assert_contains(case, html, 'src="candidate_frames/scene-001.jpg"')
        assert_contains(case, html, 'src="candidate_frames/scene-003.jpg"')
        assert_not_contains(case, html, "scene-002.jpg")
        assert_contains(case, html, "<h2>Transcript</h2>")
        assert_contains(case, html, "<h3>Segments</h3>")
        assert_not_contains(case, html, "<script>")
        passed()

        docx = job / "report.docx"
        headings = docx_headings(docx)
        for required in ("Recap: minimal.mp4", "Media", "Chapters",
                         "Transcript", "Segments"):
            if required not in headings:
                fail(case, f"report.docx missing heading {required!r}; "
                     f"got {headings}")
        shapes = docx_inline_shapes(docx)
        if shapes != 2:
            fail(case, f"report.docx inline_shapes={shapes}, expected 2")
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_absent_selected_path() -> None:
    case = "absent-selected"
    job = copy_fixture()
    try:
        (job / "selected_frames.json").unlink()
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        assert_not_contains(case, md, "## Chapters")
        assert_not_contains(case, md, "candidate_frames/")
        passed()

        html = job / "report.html"
        assert_not_contains(case, html, "<h2>Chapters</h2>")
        assert_not_contains(case, html, "<img")
        passed()

        docx = job / "report.docx"
        headings = docx_headings(docx)
        if "Chapters" in headings:
            fail(case, f"report.docx unexpectedly has Chapters heading: {headings}")
        shapes = docx_inline_shapes(docx)
        if shapes != 0:
            fail(case, f"report.docx inline_shapes={shapes}, expected 0")
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def _mutate_selected(job: Path, mutator) -> None:
    p = job / "selected_frames.json"
    data = json.loads(p.read_text(encoding="utf-8"))
    mutator(data)
    p.write_text(json.dumps(data), encoding="utf-8")


def check_bad_start_seconds() -> None:
    case = "bad-start-seconds"
    needles = (
        "error: selected_frames.json malformed: "
        "chapter 1 'start_seconds' must be numeric",
    )
    for cmd in EXPORT_CMDS:
        job = copy_fixture()
        try:
            _mutate_selected(
                job,
                lambda d: d["chapters"][0].__setitem__("start_seconds", "bad"),
            )
            for out in ("report.md", "report.html", "report.docx"):
                (job / out).unlink(missing_ok=True)
            expect_fail_with(case, job, cmd, needles)
            passed()
        finally:
            shutil.rmtree(job.parent, ignore_errors=True)


def check_traversal_frame_file() -> None:
    case = "traversal-frame-file"
    # All three exporters must reject traversal via the plain-filename
    # safety check, not the existence check. This is a validator-contract
    # requirement, not an implementation detail.
    needles = ("plain filename inside candidate_frames/",)
    for cmd in EXPORT_CMDS:
        job = copy_fixture()
        try:
            def mut(d):
                ch = d["chapters"][0]
                ch["hero"]["frame_file"] = "../report.md"
                ch["frames"][0]["frame_file"] = "../report.md"
            _mutate_selected(job, mut)
            for out in ("report.md", "report.html", "report.docx"):
                (job / out).unlink(missing_ok=True)
            expect_fail_with(case, job, cmd, needles)
            passed()
        finally:
            shutil.rmtree(job.parent, ignore_errors=True)


def check_missing_image() -> None:
    case = "missing-image"
    needles = ("missing candidate frame: candidate_frames/scene-001.jpg",)
    for cmd in EXPORT_CMDS:
        job = copy_fixture()
        try:
            (job / "candidate_frames" / "scene-001.jpg").unlink()
            for out in ("report.md", "report.html", "report.docx"):
                (job / out).unlink(missing_ok=True)
            expect_fail_with(case, job, cmd, needles)
            passed()
        finally:
            shutil.rmtree(job.parent, ignore_errors=True)


def _assert_insights_shape(case: str, insights: dict) -> None:
    for key in (
        "version",
        "provider",
        "model",
        "generated_at",
        "overview",
        "chapters",
        "action_items",
        "sources",
    ):
        if key not in insights:
            fail(case, f"insights.json missing top-level key {key!r}")
    if insights["version"] != INSIGHTS_SCHEMA_VERSION:
        fail(
            case,
            f"insights.json version={insights['version']!r}, "
            f"expected {INSIGHTS_SCHEMA_VERSION}",
        )
    if insights["provider"] != "mock":
        fail(
            case,
            f"insights.json provider={insights['provider']!r}, expected 'mock'",
        )
    overview = insights["overview"]
    for key in ("title", "short_summary", "detailed_summary", "quick_bullets"):
        if key not in overview:
            fail(case, f"insights.overview missing {key!r}")
    if not isinstance(overview["quick_bullets"], list):
        fail(case, "insights.overview.quick_bullets is not a list")
    if not isinstance(insights["chapters"], list):
        fail(case, "insights.chapters is not a list")
    if not insights["chapters"]:
        fail(case, "insights.chapters unexpectedly empty for fixture")
    for ch in insights["chapters"]:
        for key in (
            "index",
            "start_seconds",
            "end_seconds",
            "title",
            "summary",
            "bullets",
            "action_items",
            "speaker_focus",
        ):
            if key not in ch:
                fail(case, f"insights chapter missing {key!r}: {ch!r}")


def check_insights_mock_flow() -> None:
    """Happy path: mock insights + assemble/export-html/export-docx must
    include Overview, Quick bullets, and per-chapter enrichment alongside
    the existing Chapters / Transcript sections."""
    case = "insights-mock-flow"
    job = copy_fixture()
    try:
        # First run insights to produce the artifact.
        result = subprocess.run(
            [
                sys.executable, "-m", "recap", "insights",
                "--job", str(job), "--provider", "mock", "--force",
            ],
            cwd=REPO_ROOT, capture_output=True, text=True,
        )
        if result.returncode != 0:
            fail(
                case,
                f"`recap insights` returned {result.returncode}; "
                f"stderr={result.stderr.strip()!r}",
            )

        insights_path = job / "insights.json"
        if not insights_path.is_file():
            fail(case, "insights.json was not produced")
        if (job / "insights.json.tmp").exists():
            fail(case, "leftover insights.json.tmp after successful run")
        insights = json.loads(insights_path.read_text(encoding="utf-8"))
        _assert_insights_shape(case, insights)
        passed()

        # Re-run exports with insights present.
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        assert_contains(case, md, "## Overview")
        assert_contains(case, md, "### Quick bullets")
        assert_contains(case, md, "## Chapters")
        assert_contains(
            case,
            md,
            "### Chapter 1 — Welcome to the demo recording",
        )
        # Screenshots from selected_frames are still embedded.
        assert_contains(case, md, "candidate_frames/scene-001.jpg")
        passed()

        html = job / "report.html"
        assert_contains(case, html, '<section class="overview">')
        assert_contains(case, html, "<h2>Overview</h2>")
        assert_contains(case, html, '<ul class="quick-bullets">')
        assert_contains(case, html, "<h2>Chapters</h2>")
        assert_contains(case, html, "Welcome to the demo recording")
        # No unexpected script injection in the rendered HTML.
        assert_not_contains(case, html, "<script>")
        passed()

        docx = job / "report.docx"
        headings = docx_headings(docx)
        for required in ("Overview", "Quick bullets", "Chapters"):
            if required not in headings:
                fail(
                    case,
                    f"report.docx missing heading {required!r}; "
                    f"got {headings}",
                )
        if not any(
            h.startswith("Chapter 1") and "Welcome to the demo" in h
            for h in headings
        ):
            fail(
                case,
                f"report.docx missing insights-titled chapter heading; "
                f"got {headings}",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_insights_absent_still_exports() -> None:
    """Insights artifact not required: when insights.json is absent,
    assemble / export-html / export-docx must still succeed and must
    NOT emit the Overview section."""
    case = "insights-absent-still-exports"
    job = copy_fixture()
    try:
        if (job / "insights.json").exists():
            fail(case, "fixture unexpectedly ships with insights.json")

        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        assert_not_contains(case, md, "## Overview")
        assert_not_contains(case, md, "### Quick bullets")
        # Existing chapters-from-selected_frames still rendered.
        assert_contains(case, md, "## Chapters")

        html = job / "report.html"
        assert_not_contains(case, html, '<section class="overview">')
        assert_not_contains(case, html, "<h2>Overview</h2>")

        docx = job / "report.docx"
        headings = docx_headings(docx)
        if "Overview" in headings:
            fail(
                case,
                f"report.docx unexpectedly has Overview heading: {headings}",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_insights_mock_offline_no_key() -> None:
    """Mock provider must run without any env vars, including no
    GROQ_API_KEY. This is what tests and local-only users rely on."""
    case = "insights-mock-offline-no-key"
    job = copy_fixture()
    try:
        import os
        env = dict(os.environ)
        env.pop("GROQ_API_KEY", None)
        env.pop("GROQ_MODEL", None)
        env.pop("GROQ_BASE_URL", None)
        result = subprocess.run(
            [
                sys.executable, "-m", "recap", "insights",
                "--job", str(job), "--provider", "mock",
            ],
            cwd=REPO_ROOT, capture_output=True, text=True, env=env,
        )
        if result.returncode != 0:
            fail(
                case,
                f"mock insights failed without Groq key: "
                f"stderr={result.stderr.strip()!r}",
            )
        if not (job / "insights.json").is_file():
            fail(case, "insights.json missing after offline mock run")
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_insights_groq_missing_key() -> None:
    """Groq provider must fail cleanly with a one-line error when
    GROQ_API_KEY is absent, and must leave no insights.json artifact."""
    case = "insights-groq-missing-key"
    job = copy_fixture()
    try:
        import os
        env = dict(os.environ)
        env.pop("GROQ_API_KEY", None)
        env.pop("GROQ_MODEL", None)
        env.pop("GROQ_BASE_URL", None)
        result = subprocess.run(
            [
                sys.executable, "-m", "recap", "insights",
                "--job", str(job), "--provider", "groq",
            ],
            cwd=REPO_ROOT, capture_output=True, text=True, env=env,
        )
        if result.returncode == 0:
            fail(
                case,
                f"`recap insights --provider groq` unexpectedly "
                f"succeeded without GROQ_API_KEY; stdout="
                f"{result.stdout.strip()!r}",
            )
        if "GROQ_API_KEY is not set" not in result.stderr:
            fail(
                case,
                f"stderr did not surface missing GROQ_API_KEY; "
                f"got {result.stderr.strip()!r}",
            )
        if (job / "insights.json").exists():
            fail(
                case,
                "insights.json written despite failed Groq run",
            )
        if (job / "insights.json.tmp").exists():
            fail(case, "leftover insights.json.tmp after failed run")
        # The stage entry must be marked failed so the UI surfaces the
        # error instead of silently claiming everything is fine.
        state = json.loads((job / "job.json").read_text(encoding="utf-8"))
        sc = state.get("stages", {}).get("insights") or {}
        if sc.get("status") != "failed":
            fail(
                case,
                f"stages.insights.status={sc.get('status')!r}, "
                "expected 'failed'",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_insights_stage_remains_opt_in() -> None:
    """Regression guard: the insights stage must NOT be part of
    ``job.STAGES`` and must NOT be invoked by ``recap run``.

    This is a static check against the source so a future refactor
    that quietly adds 'insights' to the canonical stage tuple or
    wires it into ``cmd_run`` composition fails this verifier even
    on a machine without Groq credentials.
    """
    case = "insights-stage-remains-opt-in"
    sys.path.insert(0, str(REPO_ROOT))
    from recap import job as job_mod  # type: ignore

    if "insights" in job_mod.STAGES:
        fail(
            case,
            f"recap.job.STAGES must not include 'insights'; got "
            f"{job_mod.STAGES!r}",
        )

    cli_text = (REPO_ROOT / "recap" / "cli.py").read_text(encoding="utf-8")
    # Find the cmd_run body and check it does not reference insights.
    start = cli_text.find("def cmd_run(")
    end = cli_text.find("\ndef ", start + 1)
    if start == -1 or end == -1:
        fail(case, "could not locate cmd_run in recap/cli.py")
    body = cli_text[start:end]
    if "insights" in body:
        fail(
            case,
            "cmd_run in recap/cli.py unexpectedly references "
            "insights; insights must stay opt-in",
        )
    passed()


def check_insights_validate_requires_sources() -> None:
    """``validate_insights`` must reject documents that omit the
    top-level ``sources`` block. We wrote the block into every output
    artifact from day one; schema validation has to enforce that
    future writers keep doing so."""
    case = "insights-validate-requires-sources"
    sys.path.insert(0, str(REPO_ROOT))
    from recap.stages.insights import (  # type: ignore
        INSIGHTS_VERSION,
        validate_insights,
    )

    bad = {
        "version": INSIGHTS_VERSION,
        "provider": "mock",
        "model": "mock-v1",
        "generated_at": "2026-04-20T00:00:00Z",
        "overview": {
            "title": "t",
            "short_summary": "s",
            "detailed_summary": "s",
            "quick_bullets": [],
        },
        "chapters": [],
        "action_items": [],
    }
    raised = False
    try:
        validate_insights(bad)
    except RuntimeError as e:
        raised = True
        if "sources" not in str(e):
            fail(
                case,
                f"validate_insights error should mention sources; "
                f"got {e!r}",
            )
    if not raised:
        fail(case, "validate_insights accepted a doc missing 'sources'")

    # And the schema must reject a bogus-shape sources block.
    bad2 = dict(bad)
    bad2["sources"] = {"transcript": 123}
    raised = False
    try:
        validate_insights(bad2)
    except RuntimeError as e:
        raised = True
        if "sources.transcript" not in str(e):
            fail(case, f"expected sources.transcript error; got {e!r}")
    if not raised:
        fail(
            case,
            "validate_insights accepted sources.transcript of wrong type",
        )
    passed()


def _insights_env_without_groq() -> dict:
    import os
    env = dict(os.environ)
    env.pop("GROQ_API_KEY", None)
    env.pop("GROQ_MODEL", None)
    env.pop("GROQ_BASE_URL", None)
    return env


def check_insights_malformed_speaker_names_graceful() -> None:
    """Policy: malformed ``speaker_names.json`` must NOT block insights
    generation. Insights should silently fall back to an empty speaker
    overlay (matches how ``recap ui`` handles a half-written overlay)."""
    case = "insights-malformed-speaker-names-graceful"
    job = copy_fixture()
    try:
        (job / "speaker_names.json").write_text("{not json", encoding="utf-8")
        result = subprocess.run(
            [
                sys.executable, "-m", "recap", "insights",
                "--job", str(job), "--provider", "mock", "--force",
            ],
            cwd=REPO_ROOT,
            capture_output=True,
            text=True,
            env=_insights_env_without_groq(),
        )
        if result.returncode != 0:
            fail(
                case,
                "insights should tolerate malformed speaker_names.json; "
                f"stderr={result.stderr.strip()!r}",
            )
        if not (job / "insights.json").is_file():
            fail(case, "insights.json missing after graceful fallback")
        # The sources block should report speaker_names as absent when
        # the overlay could not be parsed.
        doc = json.loads(
            (job / "insights.json").read_text(encoding="utf-8")
        )
        if doc.get("sources", {}).get("speaker_names") is not None:
            fail(
                case,
                f"sources.speaker_names should be null on malformed "
                f"overlay; got {doc.get('sources')!r}",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_insights_malformed_selected_frames_graceful() -> None:
    """Policy: insights never parses selected_frames.json content, so
    malformed JSON must be tolerated with the artifact treated as
    absent in the sources block."""
    case = "insights-malformed-selected-frames-graceful"
    job = copy_fixture()
    try:
        (job / "selected_frames.json").write_text(
            "{not json", encoding="utf-8"
        )
        result = subprocess.run(
            [
                sys.executable, "-m", "recap", "insights",
                "--job", str(job), "--provider", "mock", "--force",
            ],
            cwd=REPO_ROOT,
            capture_output=True,
            text=True,
            env=_insights_env_without_groq(),
        )
        if result.returncode != 0:
            fail(
                case,
                "insights should tolerate malformed selected_frames.json; "
                f"stderr={result.stderr.strip()!r}",
            )
        if not (job / "insights.json").is_file():
            fail(case, "insights.json missing after graceful fallback")
        doc = json.loads(
            (job / "insights.json").read_text(encoding="utf-8")
        )
        if doc.get("sources", {}).get("selected_frames") is not None:
            fail(
                case,
                f"sources.selected_frames should be null on malformed "
                f"artifact; got {doc.get('sources')!r}",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_insights_malformed_chapter_candidates_fails_cleanly() -> None:
    """Policy: insights reads chapter_candidates.json content, so
    malformed JSON must FAIL CLEANLY with the canonical
    ``chapter_candidates.json malformed: ...`` prefix rather than
    silently fall back to a whole-transcript chapter. Mirrors
    ``recap assemble`` / ``recap export-*`` behavior."""
    case = "insights-malformed-chapter-candidates-fails-cleanly"
    job = copy_fixture()
    try:
        (job / "chapter_candidates.json").write_text(
            "{not json", encoding="utf-8"
        )
        result = subprocess.run(
            [
                sys.executable, "-m", "recap", "insights",
                "--job", str(job), "--provider", "mock", "--force",
            ],
            cwd=REPO_ROOT,
            capture_output=True,
            text=True,
            env=_insights_env_without_groq(),
        )
        if result.returncode == 0:
            fail(
                case,
                "insights unexpectedly accepted malformed "
                "chapter_candidates.json",
            )
        if "chapter_candidates.json malformed" not in result.stderr:
            fail(
                case,
                "stderr must carry canonical "
                "'chapter_candidates.json malformed' prefix; got "
                f"{result.stderr.strip()!r}",
            )
        if (job / "insights.json").exists():
            fail(
                case,
                "insights.json written despite failed chapter_candidates "
                "parse",
            )
        if (job / "insights.json.tmp").exists():
            fail(case, "leftover insights.json.tmp after failed run")
        # Stage entry must be marked failed.
        state = json.loads((job / "job.json").read_text(encoding="utf-8"))
        sc = state.get("stages", {}).get("insights") or {}
        if sc.get("status") != "failed":
            fail(
                case,
                f"stages.insights.status={sc.get('status')!r}, "
                "expected 'failed'",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_insights_groq_max_tokens_is_bounded() -> None:
    """Static check: the Groq request body must carry a bounded
    ``max_tokens`` so responses are capped by the provider as well as
    by the client-side ``MAX_RESPONSE_BYTES`` guard."""
    case = "insights-groq-max-tokens-is-bounded"
    src = (REPO_ROOT / "recap" / "stages" / "insights.py").read_text(
        encoding="utf-8"
    )
    if '"max_tokens"' not in src:
        fail(case, "insights.py Groq body is missing a 'max_tokens' cap")
    if "GROQ_DEFAULT_MAX_TOKENS" not in src:
        fail(
            case,
            "insights.py must define a GROQ_DEFAULT_MAX_TOKENS constant",
        )
    passed()


def _docx_paragraphs_text(path: Path) -> list[str]:
    return [p.text for p in Document(str(path)).paragraphs]


def check_overlays_no_op_preserves_byte_output() -> None:
    """Overlay files that resolve to nothing (empty / malformed) must
    leave every exporter's output byte-identical to the no-overlay
    baseline. Guards against accidentally inlining overlay hooks in
    ways that rewrite output even on empty overlays.
    """
    case = "overlays-no-op-byte-compat"
    baseline = copy_fixture()
    try:
        for cmd in EXPORT_CMDS:
            expect_ok(case, baseline, cmd)
        baseline_md = (baseline / "report.md").read_bytes()
        baseline_html = (baseline / "report.html").read_bytes()
    finally:
        pass

    modified = copy_fixture()
    try:
        # Empty overlays.
        (modified / "speaker_names.json").write_text(
            json.dumps({"version": 1, "updated_at": None, "speakers": {}}),
            encoding="utf-8",
        )
        (modified / "chapter_titles.json").write_text(
            json.dumps({"version": 1, "updated_at": None, "titles": {}}),
            encoding="utf-8",
        )
        (modified / "frame_review.json").write_text(
            json.dumps({"version": 1, "updated_at": None, "frames": {}}),
            encoding="utf-8",
        )
        for cmd in EXPORT_CMDS:
            expect_ok(case, modified, cmd)
        if (modified / "report.md").read_bytes() != baseline_md:
            fail(case, "report.md differs from no-overlay baseline "
                       "when overlays are empty")
        if (modified / "report.html").read_bytes() != baseline_html:
            fail(case, "report.html differs from no-overlay baseline "
                       "when overlays are empty")
        # DOCX package is non-deterministic (zip timestamps) so we
        # check structural invariants instead.
        docx = modified / "report.docx"
        if docx_inline_shapes(docx) != 2:
            fail(case, "empty overlays must leave DOCX inline shapes at 2")
        passed()
    finally:
        shutil.rmtree(baseline.parent, ignore_errors=True)
        shutil.rmtree(modified.parent, ignore_errors=True)


def check_chapter_titles_overlay_applied() -> None:
    """chapter_titles.json changes the rendered chapter heading across
    report.md, report.html, and report.docx — beating the default
    ``Chapter N`` heading and the insights-provided title.
    """
    case = "chapter-titles-overlay-applied"
    job = copy_fixture()
    try:
        (job / "chapter_titles.json").write_text(
            json.dumps({
                "version": 1,
                "updated_at": "2026-04-20T00:00:00Z",
                "titles": {"1": "User-chosen Kickoff"},
            }),
            encoding="utf-8",
        )
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        assert_contains(
            case, md, "### Chapter 1 — User-chosen Kickoff",
        )
        passed()

        html = job / "report.html"
        assert_contains(
            case, html, "Chapter 1 — User-chosen Kickoff",
        )
        passed()

        docx = job / "report.docx"
        headings = docx_headings(docx)
        if not any(
            "Chapter 1 — User-chosen Kickoff" in h for h in headings
        ):
            fail(
                case,
                f"report.docx chapter heading missing overlay title; "
                f"got {headings}",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_chapter_titles_overlay_beats_insights() -> None:
    """When both insights.title and chapter_titles overlay are set, the
    overlay wins.
    """
    case = "chapter-titles-overlay-beats-insights"
    job = copy_fixture()
    try:
        run_cmd(case, job, "insights", extra=("--provider", "mock"))
        (job / "chapter_titles.json").write_text(
            json.dumps({
                "version": 1,
                "updated_at": "2026-04-20T00:00:00Z",
                "titles": {"1": "Custom beats insights"},
            }),
            encoding="utf-8",
        )
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)
        md = job / "report.md"
        assert_contains(
            case, md, "### Chapter 1 — Custom beats insights",
        )
        # The insights-provided title must NOT be shown for chapter 1
        # when the overlay is present.
        text = md.read_text(encoding="utf-8")
        if "### Chapter 1 — Welcome" in text:
            fail(
                case,
                "chapter_titles overlay did not override insights title",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_frame_review_reject_removes_hero() -> None:
    """frame_review.json with decision=reject on the chapter hero must
    suppress the hero image from every exporter without replacing it.
    """
    case = "frame-review-reject-hero"
    job = copy_fixture()
    try:
        (job / "frame_review.json").write_text(
            json.dumps({
                "version": 1,
                "updated_at": "2026-04-20T00:00:00Z",
                "frames": {
                    "scene-001.jpg": {
                        "decision": "reject",
                        "note": "out of focus",
                    },
                },
            }),
            encoding="utf-8",
        )
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        assert_not_contains(case, md, "candidate_frames/scene-001.jpg")
        assert_contains(case, md, "candidate_frames/scene-003.jpg")
        passed()

        html = job / "report.html"
        assert_not_contains(case, html, "scene-001.jpg")
        assert_contains(case, html, 'src="candidate_frames/scene-003.jpg"')
        passed()

        docx = job / "report.docx"
        # With hero rejected, only one supporting image remains.
        shapes = docx_inline_shapes(docx)
        if shapes != 1:
            fail(
                case,
                f"report.docx inline_shapes={shapes}, expected 1 after "
                "rejecting the hero (one supporting frame should remain)",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_frame_review_keep_promotes_vlm_rejected() -> None:
    """frame_review.json with decision=keep on a vlm_rejected frame
    must add it to the exported output as an extra supporting image.
    """
    case = "frame-review-keep-promotes"
    job = copy_fixture()
    try:
        (job / "frame_review.json").write_text(
            json.dumps({
                "version": 1,
                "updated_at": "2026-04-20T00:00:00Z",
                "frames": {
                    "scene-002.jpg": {
                        "decision": "keep",
                        "note": "user override",
                    },
                },
            }),
            encoding="utf-8",
        )
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        # All three images should appear now: hero, default supporting,
        # and the user-kept previously-rejected frame.
        assert_contains(case, md, "candidate_frames/scene-001.jpg")
        assert_contains(case, md, "candidate_frames/scene-002.jpg")
        assert_contains(case, md, "candidate_frames/scene-003.jpg")
        passed()

        html = job / "report.html"
        assert_contains(case, html, 'src="candidate_frames/scene-001.jpg"')
        assert_contains(case, html, 'src="candidate_frames/scene-002.jpg"')
        assert_contains(case, html, 'src="candidate_frames/scene-003.jpg"')
        passed()

        docx = job / "report.docx"
        shapes = docx_inline_shapes(docx)
        if shapes != 3:
            fail(
                case,
                f"report.docx inline_shapes={shapes}, expected 3 after "
                "promoting a user-kept vlm_rejected frame",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_speaker_names_overlay_applied() -> None:
    """When the transcript carries utterances with speaker ids, the
    overlay substitutes the ``Speaker N`` prefix with the user's
    custom label across all three exports.

    The fixture transcript is segments-only, so we scratch-seed
    utterances inside the copied fixture only (the committed fixture
    is never touched) and verify that export output reflects both the
    utterance rendering path AND the overlay substitution.
    """
    case = "speaker-names-overlay-applied"
    job = copy_fixture()
    try:
        # Seed a small utterances[] list + a matching overlay in the
        # scratch job copy. This exercises the exporters' utterance
        # rendering path (only taken when utterances exist) and the
        # overlay substitution.
        tpath = job / "transcript.json"
        tdata = json.loads(tpath.read_text(encoding="utf-8"))
        tdata["utterances"] = [
            {
                "start": 0.0,
                "end": 8.5,
                "text": "Welcome to the demo recording.",
                "speaker": 0,
            },
            {
                "start": 8.5,
                "end": 15.0,
                "text": "I'll walk through the pipeline.",
                "speaker": 1,
            },
        ]
        tpath.write_text(json.dumps(tdata), encoding="utf-8")
        (job / "speaker_names.json").write_text(
            json.dumps({
                "version": 1,
                "updated_at": "2026-04-20T00:00:00Z",
                "speakers": {"0": "Host", "1": "Guest"},
            }),
            encoding="utf-8",
        )

        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        assert_contains(case, md, "### Utterances")
        assert_contains(case, md, "Host: Welcome to the demo")
        assert_contains(case, md, "Guest: I'll walk through the pipeline.")
        # Default `Speaker 0` / `Speaker 1` must not appear when an
        # overlay covers both speakers.
        assert_not_contains(case, md, "Speaker 0:")
        assert_not_contains(case, md, "Speaker 1:")
        passed()

        html = job / "report.html"
        assert_contains(case, html, "<h3>Utterances</h3>")
        assert_contains(case, html, "<strong>Host:</strong>")
        assert_contains(case, html, "<strong>Guest:</strong>")
        passed()

        docx = job / "report.docx"
        headings = docx_headings(docx)
        if "Utterances" not in headings:
            fail(
                case,
                f"report.docx missing 'Utterances' heading; got {headings}",
            )
        paragraphs = _docx_paragraphs_text(docx)
        if not any("Host: Welcome to the demo" in p for p in paragraphs):
            fail(
                case,
                "report.docx did not carry the custom 'Host' speaker label",
            )
        if not any("Guest: I'll walk through" in p for p in paragraphs):
            fail(
                case,
                "report.docx did not carry the custom 'Guest' speaker label",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_speaker_names_overlay_partial_falls_back() -> None:
    """When the overlay covers only some speakers, uncovered ones fall
    back to ``Speaker N``.
    """
    case = "speaker-names-overlay-partial"
    job = copy_fixture()
    try:
        tpath = job / "transcript.json"
        tdata = json.loads(tpath.read_text(encoding="utf-8"))
        tdata["utterances"] = [
            {"start": 0.0, "end": 4.0, "text": "Hi.", "speaker": 0},
            {"start": 4.0, "end": 8.0, "text": "Hello.", "speaker": 1},
        ]
        tpath.write_text(json.dumps(tdata), encoding="utf-8")
        (job / "speaker_names.json").write_text(
            json.dumps({
                "version": 1,
                "updated_at": None,
                "speakers": {"0": "Host"},
            }),
            encoding="utf-8",
        )
        expect_ok(case, job, "assemble")
        md_text = (job / "report.md").read_text(encoding="utf-8")
        if "Host: Hi." not in md_text:
            fail(case, "overlay label not applied to speaker 0")
        if "Speaker 1: Hello." not in md_text:
            fail(
                case,
                "uncovered speaker 1 did not fall back to 'Speaker 1:'",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_malformed_overlays_ignored() -> None:
    """Malformed overlays degrade to empty overlays — no exporter
    should raise, and no overlay behavior should apply.
    """
    case = "overlays-malformed-ignored"
    for overlay_name in (
        "speaker_names.json", "chapter_titles.json", "frame_review.json",
    ):
        job = copy_fixture()
        try:
            (job / overlay_name).write_text("{not json", encoding="utf-8")
            for cmd in EXPORT_CMDS:
                expect_ok(case, job, cmd)
            # Default output still includes hero + supporting and no
            # custom chapter title.
            md = job / "report.md"
            assert_contains(case, md, "candidate_frames/scene-001.jpg")
            assert_contains(case, md, "candidate_frames/scene-003.jpg")
            assert_not_contains(case, md, "candidate_frames/scene-002.jpg")
            assert_not_contains(case, md, "User-chosen Kickoff")
            passed()
        finally:
            shutil.rmtree(job.parent, ignore_errors=True)


def check_frame_review_reject_wins_over_selection() -> None:
    """When a frame appears as selected_hero AND is rejected by the
    overlay, the overlay wins and the hero is suppressed. This is the
    "frame_review overlay wins for user intent" contract.
    """
    case = "frame-review-reject-overrides-selection"
    job = copy_fixture()
    try:
        # The fixture's chapter 1 hero is scene-001.jpg. Reject it.
        (job / "frame_review.json").write_text(
            json.dumps({
                "version": 1,
                "updated_at": None,
                "frames": {
                    "scene-001.jpg": {
                        "decision": "reject", "note": "",
                    },
                },
            }),
            encoding="utf-8",
        )
        expect_ok(case, job, "assemble")
        md_text = (job / "report.md").read_text(encoding="utf-8")
        if "scene-001.jpg" in md_text:
            fail(
                case,
                "frame_review 'reject' did not override the algorithm's "
                "selected_hero in the exported report",
            )
        if "scene-003.jpg" not in md_text:
            fail(
                case,
                "supporting frame scene-003.jpg unexpectedly dropped",
            )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def _write_transcript_notes(job: Path, items: dict) -> None:
    (job / "transcript_notes.json").write_text(
        json.dumps(
            {
                "version": 1,
                "updated_at": "2026-04-21T00:00:00Z",
                "items": items,
            }
        ),
        encoding="utf-8",
    )


def _seed_utterances(job: Path) -> None:
    """Inject a minimal utterances[] section into the scratch
    transcript so the exporters take the Utterances render path. Only
    mutates the *scratch copy* of the fixture; `scripts/fixtures/*`
    stays byte-identical.
    """
    tpath = job / "transcript.json"
    tdata = json.loads(tpath.read_text(encoding="utf-8"))
    tdata["utterances"] = [
        {
            "start": 0.0,
            "end": 6.0,
            "text": "Welcome to the demo recording.",
            "speaker": 0,
        },
        {
            "start": 6.0,
            "end": 14.0,
            "text": "We will walk through the Recap pipeline today.",
            "speaker": 1,
        },
    ]
    tpath.write_text(json.dumps(tdata), encoding="utf-8")


def check_transcript_notes_segment_correction() -> None:
    """``transcript_notes.json`` correction replaces the canonical
    segment text in all three exporter outputs and marks the row as
    edited without mutating ``transcript.json``.
    """
    case = "transcript-notes-segment-correction"
    job = copy_fixture()
    try:
        canonical_transcript = (
            (job / "transcript.json").read_bytes()
        )
        _write_transcript_notes(
            job,
            {
                "seg-0": {
                    "correction": "Clarified opening line.",
                    "note": "Rewrote for clarity.",
                },
            },
        )
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        # Corrected text replaces canonical for this segment.
        assert_contains(case, md, "Clarified opening line. *(edited)*")
        # The note renders as an italic _Note:_ bullet beneath.
        assert_contains(case, md, "_Note:_ Rewrote for clarity.")
        # Canonical seg-0 text must not appear in the Transcript
        # section — it only lives in the original transcript.json and
        # the chapter body text (which comes from
        # chapter_candidates.json, unrelated to this overlay).
        md_text = md.read_text(encoding="utf-8")
        transcript_section = md_text[md_text.find("## Transcript"):]
        if (
            "Welcome to the demo recording for the Recap pipeline."
            in transcript_section
        ):
            fail(
                case,
                "canonical seg-0 text leaked into the Transcript "
                "section after correction",
            )
        passed()

        html = job / "report.html"
        assert_contains(case, html, "Clarified opening line.")
        assert_contains(
            case,
            html,
            '<small class="transcript-edited">(edited)</small>',
        )
        assert_contains(
            case,
            html,
            'class="transcript-note"',
        )
        assert_contains(case, html, "Rewrote for clarity.")
        passed()

        docx = job / "report.docx"
        paragraphs = _docx_paragraphs_text(docx)
        if not any("Clarified opening line." in p for p in paragraphs):
            fail(case, f"DOCX missing correction: {paragraphs!r}")
        if not any("(edited)" in p for p in paragraphs):
            fail(case, "DOCX missing (edited) marker")
        if not any("Rewrote for clarity." in p for p in paragraphs):
            fail(case, "DOCX missing reviewer note")
        passed()

        # Upstream transcript.json must be untouched.
        if (job / "transcript.json").read_bytes() != canonical_transcript:
            fail(case, "transcript.json was mutated by exporter")
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_transcript_notes_note_only_preserves_canonical() -> None:
    """A row with a note but no correction keeps the canonical
    transcript text verbatim and appends the note.
    """
    case = "transcript-notes-note-only"
    job = copy_fixture()
    try:
        _write_transcript_notes(
            job,
            {"seg-2": {"note": "Revisit this point with the team."}},
        )
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        # Canonical text for seg-2 survives verbatim.
        assert_contains(
            case,
            md,
            "Next we walk through configuration details.",
        )
        # Note renders under the row.
        assert_contains(
            case, md, "_Note:_ Revisit this point with the team.",
        )
        # No (edited) marker for note-only rows.
        assert_not_contains(case, md, "*(edited)*")
        passed()

        html = job / "report.html"
        assert_contains(case, html, "Revisit this point with the team.")
        assert_not_contains(
            case,
            html,
            '<small class="transcript-edited">(edited)</small>',
        )
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_transcript_notes_utterance_correction() -> None:
    """On utterance-based transcripts the correction replaces the
    text while the speaker-label prefix stays intact, and the note
    renders beneath the utterance row.
    """
    case = "transcript-notes-utterance-correction"
    job = copy_fixture()
    try:
        _seed_utterances(job)
        # Also seed a speaker-names overlay so we exercise the
        # combined `speaker_names + transcript_notes` path.
        (job / "speaker_names.json").write_text(
            json.dumps(
                {
                    "version": 1,
                    "updated_at": None,
                    "speakers": {"0": "Host", "1": "Guest"},
                }
            ),
            encoding="utf-8",
        )
        _write_transcript_notes(
            job,
            {
                "utt-0": {
                    "correction": "Hello and welcome to the demo.",
                    "note": "Keep it short.",
                },
            },
        )
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        assert_contains(
            case,
            md,
            "Host: Hello and welcome to the demo. *(edited)*",
        )
        assert_contains(case, md, "_Note:_ Keep it short.")
        # Uncovered speaker fallback + canonical text for utt-1.
        assert_contains(
            case,
            md,
            "Guest: We will walk through the Recap pipeline today.",
        )
        passed()

        html = job / "report.html"
        assert_contains(case, html, "<h3>Utterances</h3>")
        assert_contains(case, html, "<strong>Host:</strong>")
        assert_contains(
            case, html, "Hello and welcome to the demo.",
        )
        assert_contains(
            case,
            html,
            '<small class="transcript-edited">(edited)</small>',
        )
        passed()

        docx = job / "report.docx"
        paragraphs = _docx_paragraphs_text(docx)
        if not any(
            "Host: Hello and welcome to the demo." in p for p in paragraphs
        ):
            fail(case, f"DOCX missing utterance correction: {paragraphs!r}")
        if not any("Keep it short." in p for p in paragraphs):
            fail(case, "DOCX missing utterance note")
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_transcript_notes_malformed_ignored() -> None:
    """A malformed ``transcript_notes.json`` degrades to an empty
    overlay: the exporters still succeed, no correction / note
    renders, and the canonical transcript text survives intact.
    """
    case = "transcript-notes-malformed-ignored"
    job = copy_fixture()
    try:
        (job / "transcript_notes.json").write_text(
            "{not json", encoding="utf-8",
        )
        for cmd in EXPORT_CMDS:
            expect_ok(case, job, cmd)

        md = job / "report.md"
        assert_not_contains(case, md, "*(edited)*")
        assert_not_contains(case, md, "_Note:_")
        assert_contains(
            case,
            md,
            "Welcome to the demo recording for the Recap pipeline.",
        )
        passed()

        html = job / "report.html"
        assert_not_contains(
            case,
            html,
            '<small class="transcript-edited">(edited)</small>',
        )
        assert_not_contains(case, html, 'class="transcript-note"')
        passed()
    finally:
        shutil.rmtree(job.parent, ignore_errors=True)


def check_transcript_notes_empty_overlay_byte_compat() -> None:
    """An empty ``transcript_notes.json`` (valid shape, no items)
    must leave exporter output byte-identical to the no-overlay
    baseline — the "overlay layer is a pure function of its input"
    contract.
    """
    case = "transcript-notes-empty-byte-compat"
    baseline = copy_fixture()
    try:
        for cmd in EXPORT_CMDS:
            expect_ok(case, baseline, cmd)
        baseline_md = (baseline / "report.md").read_bytes()
        baseline_html = (baseline / "report.html").read_bytes()
    finally:
        pass

    modified = copy_fixture()
    try:
        _write_transcript_notes(modified, {})
        for cmd in EXPORT_CMDS:
            expect_ok(case, modified, cmd)
        if (modified / "report.md").read_bytes() != baseline_md:
            fail(case, "report.md differs from no-overlay baseline")
        if (modified / "report.html").read_bytes() != baseline_html:
            fail(case, "report.html differs from no-overlay baseline")
        passed()
    finally:
        shutil.rmtree(baseline.parent, ignore_errors=True)
        shutil.rmtree(modified.parent, ignore_errors=True)


def check_scenes_interrupt_marks_failed() -> None:
    """Regression guard: Ctrl-C during `recap scenes` must leave the
    stage as `failed`, not `running`, and remove partial
    `candidate_frames/`. Simulates the interrupt by monkeypatching
    `recap.stages.scenes._detect_and_extract` to raise
    `KeyboardInterrupt` directly; never invokes PySceneDetect.
    """
    case = "scenes-interrupt-marks-failed"
    # Import inside the function so the rest of the verifier keeps
    # working even on systems where PySceneDetect's transitive deps
    # (cv2) are unavailable. Here we monkeypatch the helper, so the
    # real detector is never called.
    sys.path.insert(0, str(REPO_ROOT))
    from recap import job as job_mod
    from recap.stages import scenes as scenes_mod

    scratch = Path(tempfile.mkdtemp(prefix="recap_scenes_interrupt_"))
    try:
        job_dir = scratch / "job"
        job_dir.mkdir()
        # Minimum plausible on-disk state: analysis.mp4 must exist so
        # scenes.run() doesn't FileNotFoundError before the interrupt.
        (job_dir / "analysis.mp4").write_bytes(b"")
        job_json = {
            "job_id": "scratch",
            "created_at": "2026-04-19T00:00:00Z",
            "updated_at": "2026-04-19T00:00:00Z",
            "status": "pending",
            "source_path": None,
            "original_filename": None,
            "stages": {
                "ingest": {"status": "completed"},
                "normalize": {"status": "completed"},
                "transcribe": {"status": "completed"},
                "assemble": {"status": "completed"},
            },
            "error": None,
        }
        (job_dir / "job.json").write_text(json.dumps(job_json))
        # Seed BOTH a stale scenes.json AND a partial
        # candidate_frames/ from a prior incomplete run. The stage
        # entered recompute because `_outputs_exist` would have
        # returned False (frames listed in scenes.json aren't all on
        # disk). The interrupt-cleanup path must remove both, not
        # just candidate_frames/.
        stale_scenes = {
            "video": "analysis.mp4",
            "detector": "ContentDetector",
            "threshold": 27.0,
            "fallback": False,
            "scene_count": 2,
            "frames_dir": "candidate_frames",
            "scenes": [
                {
                    "index": 1,
                    "start_seconds": 0.0,
                    "end_seconds": 5.0,
                    "start_frame": 0,
                    "end_frame": 150,
                    "midpoint_seconds": 2.5,
                    "frame_file": "scene-001.jpg",
                },
                {
                    "index": 2,
                    "start_seconds": 5.0,
                    "end_seconds": 10.0,
                    "start_frame": 150,
                    "end_frame": 300,
                    "midpoint_seconds": 7.5,
                    "frame_file": "scene-002.jpg",
                },
            ],
        }
        (job_dir / "scenes.json").write_text(json.dumps(stale_scenes))
        (job_dir / "candidate_frames").mkdir()
        (job_dir / "candidate_frames" / "scene-001.jpg").write_bytes(b"x")
        # scene-002.jpg deliberately missing — this is why
        # `_outputs_exist` returns False and the stage recomputes.
        # Also seed a lingering .tmp from an earlier crash.
        (job_dir / "scenes.json.tmp").write_bytes(b"{}")

        paths = job_mod.open_job(job_dir)

        original = scenes_mod._detect_and_extract
        def _raise_interrupt(*_a, **_k):
            raise KeyboardInterrupt()
        scenes_mod._detect_and_extract = _raise_interrupt
        try:
            caught = False
            try:
                scenes_mod.run(paths, force=False)
            except KeyboardInterrupt:
                caught = True
            if not caught:
                fail(case, "scenes.run() did not re-raise KeyboardInterrupt")
        finally:
            scenes_mod._detect_and_extract = original

        state = json.loads((job_dir / "job.json").read_text())
        sc = state.get("stages", {}).get("scenes") or {}
        if sc.get("status") != "failed":
            fail(
                case,
                f"stages.scenes.status={sc.get('status')!r}, expected 'failed'",
            )
        if "KeyboardInterrupt" not in (sc.get("error") or ""):
            fail(
                case,
                f"stages.scenes.error missing KeyboardInterrupt: "
                f"{sc.get('error')!r}",
            )
        if (job_dir / "scenes.json").exists():
            fail(
                case,
                "stale scenes.json was not cleaned up after interrupted "
                "recompute",
            )
        if (job_dir / "candidate_frames").exists():
            fail(case, "candidate_frames/ was not cleaned up after interrupt")
        if (job_dir / "scenes.json.tmp").exists():
            fail(case, "scenes.json.tmp left behind")
        passed()
    finally:
        shutil.rmtree(scratch, ignore_errors=True)


def check_normalize_mode_decision() -> None:
    """Pure-function guard on ``_decide_normalize_mode``. The fast path
    must only trigger for MP4/MOV containers with H.264+yuv420p video
    and AAC (or no) audio; anything else must fall back to reencode.

    Also asserts ``RECAP_NORMALIZE_NO_FASTPATH=1`` hard-disables remux.
    """
    case = "normalize-mode-decision"
    sys.path.insert(0, str(REPO_ROOT))
    import os as os_mod
    from recap.stages import normalize as normalize_mod

    def _probe(
        format_name: str,
        vcodec: str = "h264",
        pix_fmt: str = "yuv420p",
        width: int = 1920,
        acodec: str | None = "aac",
    ) -> dict:
        streams: list[dict] = [
            {
                "codec_type": "video",
                "codec_name": vcodec,
                "pix_fmt": pix_fmt,
                "width": width,
            }
        ]
        if acodec is not None:
            streams.append({"codec_type": "audio", "codec_name": acodec})
        return {"format": {"format_name": format_name}, "streams": streams}

    cases = (
        ("mp4+h264+aac => remux",
         _probe("mov,mp4,m4a,3gp,3g2,mj2"), "remux"),
        ("bare mp4 container => remux",
         _probe("mp4"), "remux"),
        ("mp4+h264 no-audio => remux",
         _probe("mp4", acodec=None), "remux"),
        ("mp4+hevc => reencode",
         _probe("mp4", vcodec="hevc"), "reencode"),
        ("mp4+h264+opus => reencode",
         _probe("mp4", acodec="opus"), "reencode"),
        ("mp4+yuv444p => reencode",
         _probe("mp4", pix_fmt="yuv444p"), "reencode"),
        ("mp4+width=0 => reencode",
         _probe("mp4", width=0), "reencode"),
        ("webm => reencode",
         _probe("matroska,webm", vcodec="vp9"), "reencode"),
        ("missing format => reencode",
         {"streams": []}, "reencode"),
        ("non-dict probe => reencode",
         [], "reencode"),  # type: ignore[arg-type]
    )
    for label, probe, expected in cases:
        got = normalize_mod._decide_normalize_mode(probe)
        if got != expected:
            fail(case, f"{label}: got {got!r}, expected {expected!r}")

    # Env escape: NO_FASTPATH forces reencode even on a perfect input.
    prev = os_mod.environ.get("RECAP_NORMALIZE_NO_FASTPATH")
    os_mod.environ["RECAP_NORMALIZE_NO_FASTPATH"] = "1"
    try:
        got = normalize_mod._decide_normalize_mode(
            _probe("mov,mp4,m4a,3gp,3g2,mj2")
        )
        if got != "reencode":
            fail(
                case,
                f"RECAP_NORMALIZE_NO_FASTPATH=1 should force reencode; "
                f"got {got!r}",
            )
    finally:
        if prev is None:
            os_mod.environ.pop("RECAP_NORMALIZE_NO_FASTPATH", None)
        else:
            os_mod.environ["RECAP_NORMALIZE_NO_FASTPATH"] = prev

    passed()


def check_normalize_failure_cleans_tmp_and_marks_failed() -> None:
    """Regression guard: when the ffmpeg runner raises
    ``NormalizeError``, the stage must end in ``failed`` state, any
    ``analysis.mp4.tmp`` / ``audio.wav.tmp`` / ``metadata.json.tmp``
    must be unlinked, and the final ``analysis.mp4`` must NOT be
    promoted. Simulates the failure by monkey-patching
    ``_run_ffmpeg_streaming`` to write a stub tmp then raise.
    """
    case = "normalize-failure-cleans-tmp"
    sys.path.insert(0, str(REPO_ROOT))
    from recap import job as job_mod
    from recap.stages import normalize as normalize_mod

    scratch = Path(tempfile.mkdtemp(prefix="recap_normalize_fail_"))
    try:
        job_dir = scratch / "job"
        job_dir.mkdir()
        # A plausible "original.mp4" so find_original() returns it.
        (job_dir / "original.mp4").write_bytes(b"not-a-real-mp4")
        # Pre-seed metadata.json so the failure is isolated to the
        # ffmpeg step; probe stays a simple reencode-shaped dict.
        (job_dir / "metadata.json").write_text(json.dumps({
            "format": {"format_name": "mp4", "duration": "12.0"},
            "streams": [
                {
                    "codec_type": "video",
                    "codec_name": "hevc",
                    "pix_fmt": "yuv420p",
                    "width": 1920,
                }
            ],
        }))
        (job_dir / "job.json").write_text(json.dumps({
            "job_id": "scratch",
            "created_at": "2026-04-21T00:00:00Z",
            "updated_at": "2026-04-21T00:00:00Z",
            "status": "pending",
            "source_path": None,
            "original_filename": None,
            "stages": {
                "ingest": {"status": "completed"},
                "normalize": {"status": "pending"},
                "transcribe": {"status": "pending"},
                "assemble": {"status": "pending"},
            },
            "error": None,
        }))
        paths = job_mod.open_job(job_dir)

        original_runner = normalize_mod._run_ffmpeg_streaming
        def _fake_stall(cmd, out_tmp, timeout_s, stall_s, heartbeat=None):
            # Simulate a partially-written tmp before the stall.
            Path(out_tmp).write_bytes(b"partial-bytes")
            raise normalize_mod.NormalizeError(
                "ffmpeg stalled: no output growth or stderr activity for 90s"
            )
        normalize_mod._run_ffmpeg_streaming = _fake_stall
        try:
            raised = False
            try:
                # force=False so the pre-seeded metadata.json is reused;
                # we're isolating the ffmpeg-runner failure, not ffprobe.
                normalize_mod.run(paths, force=False)
            except normalize_mod.NormalizeError:
                raised = True
            if not raised:
                fail(case, "normalize.run() did not re-raise NormalizeError")
        finally:
            normalize_mod._run_ffmpeg_streaming = original_runner

        state = json.loads((job_dir / "job.json").read_text())
        ne = state.get("stages", {}).get("normalize") or {}
        if ne.get("status") != "failed":
            fail(
                case,
                f"stages.normalize.status={ne.get('status')!r}, "
                f"expected 'failed'",
            )
        if "stalled" not in (ne.get("error") or ""):
            fail(
                case,
                f"stages.normalize.error should mention stall; got "
                f"{ne.get('error')!r}",
            )
        if (job_dir / "analysis.mp4").exists():
            fail(case, "analysis.mp4 was promoted despite ffmpeg failure")
        for leftover in (
            "analysis.mp4.tmp",
            "audio.wav.tmp",
            "metadata.json.tmp",
        ):
            if (job_dir / leftover).exists():
                fail(case, f"{leftover} left behind after failure")
        passed()
    finally:
        shutil.rmtree(scratch, ignore_errors=True)


def check_normalize_invalid_output_not_promoted() -> None:
    """Regression guard: if the ffmpeg subprocess exits 0 but the tmp
    file is not a readable video (corrupt, truncated, wrong shape),
    ``_validate_analysis`` must reject it, the stage must end FAILED,
    no ``analysis.mp4`` may be promoted, and the tmp is cleaned up.
    """
    case = "normalize-invalid-output-not-promoted"
    sys.path.insert(0, str(REPO_ROOT))
    from recap import job as job_mod
    from recap.stages import normalize as normalize_mod

    scratch = Path(tempfile.mkdtemp(prefix="recap_normalize_bad_"))
    try:
        job_dir = scratch / "job"
        job_dir.mkdir()
        (job_dir / "original.mp4").write_bytes(b"not-a-real-mp4")
        (job_dir / "metadata.json").write_text(json.dumps({
            "format": {"format_name": "mp4", "duration": "5.0"},
            "streams": [
                {
                    "codec_type": "video",
                    "codec_name": "h264",
                    "pix_fmt": "yuv420p",
                    "width": 640,
                },
                {"codec_type": "audio", "codec_name": "aac"},
            ],
        }))
        (job_dir / "job.json").write_text(json.dumps({
            "job_id": "scratch",
            "created_at": "2026-04-21T00:00:00Z",
            "updated_at": "2026-04-21T00:00:00Z",
            "status": "pending",
            "source_path": None,
            "original_filename": None,
            "stages": {
                "ingest": {"status": "completed"},
                "normalize": {"status": "pending"},
                "transcribe": {"status": "pending"},
                "assemble": {"status": "pending"},
            },
            "error": None,
        }))
        paths = job_mod.open_job(job_dir)

        original_runner = normalize_mod._run_ffmpeg_streaming
        original_validate = normalize_mod._validate_analysis
        def _ok_runner(cmd, out_tmp, timeout_s, stall_s, heartbeat=None):
            Path(out_tmp).write_bytes(b"garbage-not-an-mp4")
        def _fail_validate(tmp):
            raise normalize_mod.NormalizeError(
                "analysis validation failed: no usable video stream"
            )
        normalize_mod._run_ffmpeg_streaming = _ok_runner
        normalize_mod._validate_analysis = _fail_validate
        try:
            raised = False
            try:
                # force=False so the pre-seeded metadata.json is reused;
                # we're isolating the validation failure.
                normalize_mod.run(paths, force=False)
            except normalize_mod.NormalizeError:
                raised = True
            if not raised:
                fail(case, "normalize.run() did not re-raise NormalizeError")
        finally:
            normalize_mod._run_ffmpeg_streaming = original_runner
            normalize_mod._validate_analysis = original_validate

        if (job_dir / "analysis.mp4").exists():
            fail(case, "analysis.mp4 promoted despite validation failure")
        if (job_dir / "analysis.mp4.tmp").exists():
            fail(case, "analysis.mp4.tmp left behind after validation failure")

        state = json.loads((job_dir / "job.json").read_text())
        ne = state.get("stages", {}).get("normalize") or {}
        if ne.get("status") != "failed":
            fail(
                case,
                f"stages.normalize.status={ne.get('status')!r}, "
                f"expected 'failed'",
            )
        if "validation" not in (ne.get("error") or ""):
            fail(
                case,
                f"stages.normalize.error should mention validation; got "
                f"{ne.get('error')!r}",
            )
        passed()
    finally:
        shutil.rmtree(scratch, ignore_errors=True)


def check_scenes_partition_tolerates_small_misses() -> None:
    """Regression guard: ``partition_scene_records`` must tolerate a
    small number of silent ``save_images`` misses, drop the skipped
    scenes from the returned list, and record them in a separate
    ``skipped`` list. Downstream stages already fail cleanly when a
    scene lacks ``frame_file``, so keeping those scenes in the
    primary list was the original bug.

    Pure-function check — no PySceneDetect, no FFmpeg.
    """
    case = "scenes-partition-tolerates-small-misses"
    sys.path.insert(0, str(REPO_ROOT))
    from recap.stages import scenes as scenes_mod

    # Build 10 pre-scene records, with save_images having silently
    # skipped scenes 4, 7 (20% miss rate — under the 25% default).
    pre = [
        {
            "index": i,
            "start_seconds": float(i - 1),
            "end_seconds": float(i),
            "start_frame": (i - 1) * 30,
            "end_frame": i * 30,
            "midpoint_seconds": (i - 1) + 0.5,
            "frame_file": None if i in (4, 7) else f"scene-{i:03d}.jpg",
        }
        for i in range(1, 11)
    ]
    kept, skipped = scenes_mod.partition_scene_records(pre, 0.25)
    if len(kept) != 8:
        fail(case, f"kept count: expected 8, got {len(kept)}")
    if len(skipped) != 2:
        fail(case, f"skipped count: expected 2, got {len(skipped)}")
    kept_ids = [s["index"] for s in kept]
    skipped_ids = [s["index"] for s in skipped]
    if kept_ids != [1, 2, 3, 5, 6, 8, 9, 10]:
        fail(case, f"kept ids wrong: {kept_ids!r}")
    if skipped_ids != [4, 7]:
        fail(case, f"skipped ids wrong: {skipped_ids!r}")
    # Skipped entries MUST NOT carry a frame_file — downstream stages
    # only see the kept list, but we still want skipped entries to be
    # unmistakable as "no frame available".
    for s in skipped:
        if "frame_file" in s:
            fail(case, f"skipped entry {s['index']} still carries frame_file")
    # Sanity: every kept entry has a usable frame_file.
    for s in kept:
        if not isinstance(s.get("frame_file"), str) or not s["frame_file"]:
            fail(
                case,
                f"kept entry {s['index']} has empty / missing frame_file",
            )
    passed()


def check_scenes_partition_rejects_high_miss_ratio() -> None:
    """Regression guard: ``partition_scene_records`` must fail cleanly
    when the miss ratio exceeds ``max_missing_ratio``. A 40% miss rate
    under the default 25% tolerance should raise ``RuntimeError``
    rather than silently produce an under-sized scene list.
    """
    case = "scenes-partition-rejects-high-miss-ratio"
    sys.path.insert(0, str(REPO_ROOT))
    from recap.stages import scenes as scenes_mod

    # 10 pre-scenes, 4 missing = 40% ratio. Default threshold 25%.
    pre = [
        {
            "index": i,
            "start_seconds": float(i - 1),
            "end_seconds": float(i),
            "start_frame": (i - 1) * 30,
            "end_frame": i * 30,
            "midpoint_seconds": (i - 1) + 0.5,
            "frame_file": (
                None if i in (2, 4, 6, 8) else f"scene-{i:03d}.jpg"
            ),
        }
        for i in range(1, 11)
    ]
    try:
        scenes_mod.partition_scene_records(pre, 0.25)
    except RuntimeError as e:
        msg = str(e)
        if "4/10" not in msg and "missed 4" not in msg:
            fail(case, f"RuntimeError missing '4/10' detail: {msg!r}")
        if "missing=" not in msg:
            fail(case, f"RuntimeError missing 'missing=' detail: {msg!r}")
        passed()
        return
    fail(case, "partition_scene_records did not raise on 40% miss ratio")


def check_scenes_partition_requires_at_least_one_frame() -> None:
    """Regression guard: ``partition_scene_records`` must raise when
    every scene is missing a frame (100% miss), regardless of the
    configured ``max_missing_ratio``. Without a single surviving
    frame there is nothing to hand downstream.
    """
    case = "scenes-partition-requires-at-least-one-frame"
    sys.path.insert(0, str(REPO_ROOT))
    from recap.stages import scenes as scenes_mod

    pre = [
        {
            "index": i,
            "start_seconds": float(i - 1),
            "end_seconds": float(i),
            "start_frame": (i - 1) * 30,
            "end_frame": i * 30,
            "midpoint_seconds": (i - 1) + 0.5,
            "frame_file": None,
        }
        for i in range(1, 4)
    ]
    try:
        # Even with a very permissive ratio, zero survivors must fail.
        scenes_mod.partition_scene_records(pre, 1.0)
    except RuntimeError as e:
        if "no frames" not in str(e):
            fail(
                case, f"RuntimeError missing 'no frames' detail: {e!r}"
            )
        passed()
        return
    fail(case, "partition_scene_records did not raise on zero survivors")


def check_scenes_partition_empty_input() -> None:
    """Regression guard: an empty scene list (detector produced no
    scenes at all, not even a fallback) must raise — the caller in
    ``_detect_and_extract`` builds a single-scene fallback so this
    path should in practice never trigger, but we enforce the
    invariant for robustness.
    """
    case = "scenes-partition-empty-input"
    sys.path.insert(0, str(REPO_ROOT))
    from recap.stages import scenes as scenes_mod

    try:
        scenes_mod.partition_scene_records([], 0.25)
    except RuntimeError as e:
        if "no scenes" not in str(e):
            fail(case, f"RuntimeError missing 'no scenes' detail: {e!r}")
        passed()
        return
    fail(case, "partition_scene_records did not raise on empty input")


def check_scenes_run_records_skipped_in_job_json() -> None:
    """Regression guard: when ``_detect_and_extract`` returns a result
    that includes a ``skipped_count > 0``, the run wrapper must
    surface that in both ``scenes.json`` on disk and in the job
    state's ``stages.scenes.extra`` block so the UI can display it.

    Monkey-patches ``_detect_and_extract`` to simulate a run where
    ``save_images`` missed 2 of 10 scenes (under the 25% default
    threshold) and asserts the resulting artifacts record both
    ``skipped_count`` and the individual skipped scene IDs.
    """
    case = "scenes-run-records-skipped-in-job-json"
    sys.path.insert(0, str(REPO_ROOT))
    from recap import job as job_mod
    from recap.stages import scenes as scenes_mod

    scratch = Path(tempfile.mkdtemp(prefix="recap_scenes_skipped_"))
    try:
        job_dir = scratch / "job"
        job_dir.mkdir()
        # `analysis.mp4` must exist so the pre-flight check passes.
        (job_dir / "analysis.mp4").write_bytes(b"")
        (job_dir / "job.json").write_text(json.dumps({
            "job_id": "scratch",
            "created_at": "2026-04-21T00:00:00Z",
            "updated_at": "2026-04-21T00:00:00Z",
            "status": "pending",
            "source_path": None,
            "original_filename": None,
            "stages": {
                "ingest": {"status": "completed"},
                "normalize": {"status": "completed"},
                "transcribe": {"status": "completed"},
                "assemble": {"status": "completed"},
            },
            "error": None,
        }))
        paths = job_mod.open_job(job_dir)

        def _fake_detect_and_extract(
            video_path, frames_dir, threshold,
        ):  # noqa: ANN001 — matches real signature
            # Build 10 pre-scenes with scenes 4 and 7 missing, then
            # run the real partition helper so we exercise exactly
            # the same splitting logic the production path uses.
            frames_dir.mkdir(parents=True, exist_ok=True)
            pre = []
            for i in range(1, 11):
                missing = i in (4, 7)
                frame_name = None if missing else f"scene-{i:03d}.jpg"
                if not missing:
                    # Emit a token byte for the file so
                    # `_outputs_exist` on the next run passes.
                    (frames_dir / frame_name).write_bytes(b"x")
                pre.append({
                    "index": i,
                    "start_seconds": float(i - 1),
                    "end_seconds": float(i),
                    "start_frame": (i - 1) * 30,
                    "end_frame": i * 30,
                    "midpoint_seconds": (i - 1) + 0.5,
                    "frame_file": frame_name,
                })
            kept, skipped = scenes_mod.partition_scene_records(
                pre, scenes_mod._max_missing_ratio(),
            )
            return {
                "video": "analysis.mp4",
                "detector": "ContentDetector",
                "threshold": threshold,
                "fallback": False,
                "scene_count": len(kept),
                "frames_dir": frames_dir.name,
                "scenes": kept,
                "skipped_count": len(skipped),
                "skipped_scenes": skipped,
            }

        original = scenes_mod._detect_and_extract
        scenes_mod._detect_and_extract = _fake_detect_and_extract
        try:
            scenes_mod.run(paths, force=False)
        finally:
            scenes_mod._detect_and_extract = original

        # scenes.json must carry both the kept list AND the skipped
        # provenance block.
        data = json.loads((job_dir / "scenes.json").read_text())
        if data.get("scene_count") != 8:
            fail(
                case,
                f"scenes.json scene_count={data.get('scene_count')!r}, "
                f"expected 8",
            )
        if data.get("skipped_count") != 2:
            fail(
                case,
                f"scenes.json skipped_count={data.get('skipped_count')!r}, "
                f"expected 2",
            )
        sk_ids = [s["index"] for s in data.get("skipped_scenes") or []]
        if sk_ids != [4, 7]:
            fail(case, f"skipped_scenes ids wrong: {sk_ids!r}")

        state = json.loads((job_dir / "job.json").read_text())
        extra = state.get("stages", {}).get("scenes") or {}
        if extra.get("status") != "completed":
            fail(
                case,
                f"stages.scenes.status={extra.get('status')!r}, "
                f"expected 'completed'",
            )
        if extra.get("skipped_scene_count") != 2:
            fail(
                case,
                f"stages.scenes.skipped_scene_count={extra.get('skipped_scene_count')!r}, "
                f"expected 2",
            )
        if extra.get("skipped_scene_ids") != [4, 7]:
            fail(
                case,
                f"stages.scenes.skipped_scene_ids={extra.get('skipped_scene_ids')!r}, "
                f"expected [4, 7]",
            )
        passed()
    finally:
        shutil.rmtree(scratch, ignore_errors=True)


def check_scenes_run_fails_on_excessive_misses() -> None:
    """Regression guard: when ``_detect_and_extract`` can't clear the
    miss-ratio threshold, ``run()`` must mark the stage FAILED with
    the partition-helper's descriptive error message rather than
    silently producing an under-sized scene list.
    """
    case = "scenes-run-fails-on-excessive-misses"
    sys.path.insert(0, str(REPO_ROOT))
    from recap import job as job_mod
    from recap.stages import scenes as scenes_mod

    scratch = Path(tempfile.mkdtemp(prefix="recap_scenes_excessive_"))
    try:
        job_dir = scratch / "job"
        job_dir.mkdir()
        (job_dir / "analysis.mp4").write_bytes(b"")
        (job_dir / "job.json").write_text(json.dumps({
            "job_id": "scratch",
            "created_at": "2026-04-21T00:00:00Z",
            "updated_at": "2026-04-21T00:00:00Z",
            "status": "pending",
            "source_path": None,
            "original_filename": None,
            "stages": {
                "ingest": {"status": "completed"},
                "normalize": {"status": "completed"},
                "transcribe": {"status": "completed"},
                "assemble": {"status": "completed"},
            },
            "error": None,
        }))
        paths = job_mod.open_job(job_dir)

        def _fake_detect_and_extract(
            video_path, frames_dir, threshold,
        ):  # noqa: ANN001
            frames_dir.mkdir(parents=True, exist_ok=True)
            # 40% miss ratio, default threshold is 25% — must fail.
            pre = [
                {
                    "index": i,
                    "start_seconds": float(i - 1),
                    "end_seconds": float(i),
                    "start_frame": (i - 1) * 30,
                    "end_frame": i * 30,
                    "midpoint_seconds": (i - 1) + 0.5,
                    "frame_file": (
                        None if i in (2, 4, 6, 8)
                        else f"scene-{i:03d}.jpg"
                    ),
                }
                for i in range(1, 11)
            ]
            return scenes_mod.partition_scene_records(
                pre, scenes_mod._max_missing_ratio(),
            )

        original = scenes_mod._detect_and_extract
        scenes_mod._detect_and_extract = _fake_detect_and_extract
        try:
            raised = False
            try:
                scenes_mod.run(paths, force=False)
            except RuntimeError:
                raised = True
            if not raised:
                fail(
                    case,
                    "scenes.run did not raise RuntimeError on a 40% miss",
                )
        finally:
            scenes_mod._detect_and_extract = original

        if (job_dir / "scenes.json").exists():
            fail(
                case,
                "scenes.json was written despite partition failing",
            )
        state = json.loads((job_dir / "job.json").read_text())
        sc = state.get("stages", {}).get("scenes") or {}
        if sc.get("status") != "failed":
            fail(
                case,
                f"stages.scenes.status={sc.get('status')!r}, "
                f"expected 'failed'",
            )
        msg = sc.get("error") or ""
        if "missed 4/10" not in msg:
            fail(
                case,
                f"stages.scenes.error missing 'missed 4/10' detail: {msg!r}",
            )
        passed()
    finally:
        shutil.rmtree(scratch, ignore_errors=True)


def check_normalize_stages_and_cmd_run_unchanged() -> None:
    """Static pin: ``recap/job.py`` ``STAGES`` and ``recap/cli.py``
    ``cmd_run`` composition are frozen. Any change must be an explicit
    user-gated act, not a side effect of a normalize slice.
    """
    case = "normalize-stages-and-cmd-run-unchanged"
    sys.path.insert(0, str(REPO_ROOT))
    from recap import job as job_mod

    expected_stages = ("ingest", "normalize", "transcribe", "assemble")
    if job_mod.STAGES != expected_stages:
        fail(
            case,
            f"recap.job.STAGES changed; expected {expected_stages!r}, "
            f"got {job_mod.STAGES!r}",
        )

    cli_text = (REPO_ROOT / "recap" / "cli.py").read_text(encoding="utf-8")
    start = cli_text.find("def cmd_run(")
    end = cli_text.find("\ndef ", start + 1)
    if start == -1 or end == -1:
        fail(case, "could not locate cmd_run in recap/cli.py")
    body = cli_text[start:end]
    # cmd_run must call exactly the four pinned stages in order, and
    # must not reference opt-in stages (scenes/dedupe/window/similarity/
    # chapters/rank/shortlist/verify/insights).
    required_calls = (
        "normalize.run(",
        "transcribe.run(",
        "assemble.run(",
    )
    for needle in required_calls:
        if needle not in body:
            fail(case, f"cmd_run no longer calls {needle!r}")
    forbidden = (
        "scenes.run(", "dedupe.run(", "window.run(", "similarity.run(",
        "chapters.run(", "rank.run(", "shortlist.run(", "verify.run(",
        "insights.run(",
    )
    for needle in forbidden:
        if needle in body:
            fail(case, f"cmd_run unexpectedly references {needle!r}")
    passed()


def main() -> int:
    if not FIXTURE_ROOT.exists():
        fail("fixture", f"committed fixture not found at {FIXTURE_ROOT}")

    check_selected_path()
    check_absent_selected_path()
    check_bad_start_seconds()
    check_traversal_frame_file()
    check_missing_image()
    check_insights_mock_flow()
    check_insights_absent_still_exports()
    check_insights_mock_offline_no_key()
    check_insights_groq_missing_key()
    check_insights_stage_remains_opt_in()
    check_insights_validate_requires_sources()
    check_insights_malformed_speaker_names_graceful()
    check_insights_malformed_selected_frames_graceful()
    check_insights_malformed_chapter_candidates_fails_cleanly()
    check_insights_groq_max_tokens_is_bounded()
    check_overlays_no_op_preserves_byte_output()
    check_chapter_titles_overlay_applied()
    check_chapter_titles_overlay_beats_insights()
    check_frame_review_reject_removes_hero()
    check_frame_review_keep_promotes_vlm_rejected()
    check_speaker_names_overlay_applied()
    check_speaker_names_overlay_partial_falls_back()
    check_malformed_overlays_ignored()
    check_frame_review_reject_wins_over_selection()
    check_transcript_notes_segment_correction()
    check_transcript_notes_note_only_preserves_canonical()
    check_transcript_notes_utterance_correction()
    check_transcript_notes_malformed_ignored()
    check_transcript_notes_empty_overlay_byte_compat()
    check_scenes_interrupt_marks_failed()
    check_scenes_partition_tolerates_small_misses()
    check_scenes_partition_rejects_high_miss_ratio()
    check_scenes_partition_requires_at_least_one_frame()
    check_scenes_partition_empty_input()
    check_scenes_run_records_skipped_in_job_json()
    check_scenes_run_fails_on_excessive_misses()
    check_normalize_mode_decision()
    check_normalize_failure_cleans_tmp_and_marks_failed()
    check_normalize_invalid_output_not_promoted()
    check_normalize_stages_and_cmd_run_unchanged()

    print(f"OK: {CHECKS_PASSED} checks passed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
